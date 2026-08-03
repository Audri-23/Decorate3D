import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_documentation():
    doc = docx.Document()

    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles & Colors
    gold_color = RGBColor(161, 122, 22) # #A17A16
    dark_color = RGBColor(30, 35, 42)    # #1E232A
    gray_color = RGBColor(100, 100, 100)

    # Document Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("Decorate3D: Smart C2C Used-Furniture & Interior Design Marketplace")
    run_title.font.name = "Georgia"
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = gold_color

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle.add_run("Project Implementation & Technical Documentation Report\n")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(14)
    run_sub.font.bold = True
    run_sub.font.color.rgb = dark_color

    # Student Details Metadata Box / Table
    table_meta = doc.add_table(rows=6, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Course", "CSE471 — System Analysis and Design (Lab Section 14, Summer 2025)"),
        ("Institution", "BRAC University"),
        ("Group Number", "Group 04"),
        ("Student Name", "Muhtasim Ahmed"),
        ("Student ID", "23101325"),
        ("Graded Feature", "Module 1 • Feature 2: Interactive 360° Rotatable 3D Canvas Renderer for Product Details")
    ]
    for i, (label, val) in enumerate(meta_data):
        row = table_meta.rows[i]
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(label)
        r0.font.bold = True
        r0.font.size = Pt(10.5)
        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.size = Pt(10.5)
        if label == "Graded Feature":
            r1.font.bold = True
            r1.font.color.rgb = gold_color

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        r = h.add_run(text)
        r.font.name = "Georgia"
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = dark_color
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        r = h.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(12.5)
        r.font.bold = True
        r.font.color.rgb = gold_color
        return h

    # Section 1
    add_heading_1("1. Executive Summary")
    p = doc.add_paragraph(
        "This document presents the detailed technical documentation for Decorate3D, a peer-to-peer (C2C) used-furniture "
        "marketplace web application developed for the CSE471 System Analysis and Design course project at BRAC University. "
        "The project is architected following a strict Unified Model-View-Controller (MVC) pattern (with dedicated top-level "
        "models/, views/, and controllers/ directories directly under the root). It specifically constructs Module 1 • Feature 2 "
        "assigned to Muhtasim Ahmed, alongside baseline marketplace capabilities (Authentication, Product Hero View, Catalog, "
        "Cart Drawer, and Student Profile) and integration hooks for team members' Module 2 and 3 features."
    )

    # Section 2
    add_heading_1("2. Tech Stack & Framework Breakdown")
    p_tech = doc.add_paragraph("The application utilizes modern industry-standard frameworks and libraries:")

    tech_table = doc.add_table(rows=7, cols=3)
    tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Layer", "Framework / Technology", "Role & Purpose"]
    hdr_cells = tech_table.rows[0].cells
    for i, h_text in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h_text)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="A17A16"/>')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)

    tech_rows = [
        ("Frontend UI Framework", "DaisyUI (v4.12) + Tailwind CSS (v3.4)", "Styled buttons, cards, badges, modals, floating toolbars, and responsive page wrappers matching screenshots."),
        ("3D Engine", "Three.js (v0.164) + WebGL", "Renders 360° rotatable 3D furniture models, studio lighting, materials, zoom/elevation sliders, and wireframe mode."),
        ("Frontend Application Core", "React.js (v18.3) + Vite (v5.4)", "Reactive user interface components, client-side routing, state management, and high-speed bundling."),
        ("Vector Icons", "Lucide React", "Icons for 3D navigation, orbit controls, zoom, layers, cart, and shield badges."),
        ("Backend REST Server", "Node.js + Express.js (v4.19)", "RESTful API server providing endpoints (/api/products, /api/auth, /api/modules)."),
        ("Database & ODM", "MongoDB + Mongoose (v8.4)", "NoSQL database schemas with seed data for offline and online execution.")
    ]

    for row_idx, data in enumerate(tech_rows, start=1):
        cells = tech_table.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            p = cells[col_idx].paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 3
    add_heading_1("3. Unified MVC System Architecture")
    p = doc.add_paragraph(
        "To strictly follow MVC architecture guidelines without nested frontend/backend directory confusion, the project "
        "is organized directly under the root workspace into top-level models/, views/, and controllers/ folders:"
    )

    mvc_structure = [
        ("models/ (Model Layer)", [
            "ProductModel.js: Mongoose schema defining furniture items, craftsmanship notes, prices, and 3D model properties (polygon counts, LOD levels, geometry type).",
            "UserModel.js: User account schema supporting Buyer, Seller, Courier, and Admin roles.",
            "OrderModel.js: Schema for transaction orders and escrow holding states.",
            "Viewer3DStore.js: State store for 3D canvas camera angles, material selections, and rotation states.",
            "seedData.js: Initial sample furniture catalog matching reference screenshots."
        ]),
        ("views/ (View Layer)", [
            "components/Navbar.jsx: DaisyUI top header navigation bar matching Screenshot 1.",
            "components/ProductCard.jsx: DaisyUI catalog cards displaying '3D Model Available' and AI condition badges.",
            "components/Viewer3DCanvas.jsx: Three.js WebGL 360-degree interactive 3D model canvas renderer.",
            "components/Viewer3DModal.jsx: Core Module-1-Feature-2 3D inspector modal matching Screenshot 2.",
            "components/RoomPlannerPreview.jsx: Spatial 3D room placement view matching Screenshot 3.",
            "components/AuthModal.jsx: DaisyUI login and registration modal.",
            "pages/ProductDetailPage.jsx: Product details hero page matching Screenshot 1 with prominent 'LAUNCH INTERACTIVE 3D INSPECTOR' button.",
            "pages/MarketplacePage.jsx: Catalog page with category tabs and AI condition filters.",
            "pages/ProfilePage.jsx: Student profile view displaying Muhtasim Ahmed's credentials."
        ]),
        ("controllers/ (Controller Layer)", [
            "productController.js: REST API handlers for catalog listings and 3D specs.",
            "authController.js: User authentication, session management, and profile handlers.",
            "useViewer3DController.js: Custom React hook and UI controller managing 360° auto-rotation, rotation speed, camera distance, elevation, and material variants.",
            "moduleHooksController.js: Reserved API integration endpoints for team members' Module 2 and 3 features."
        ])
    ]

    for folder_title, items in mvc_structure:
        add_heading_2(folder_title)
        for item in items:
            p_item = doc.add_paragraph(style='List Bullet')
            r_item = p_item.add_run(item)
            r_item.font.size = Pt(10)

    # Section 4
    add_heading_1("4. What Was Built & Feature Details")
    add_heading_2("Graded Feature: Module 1 • Feature 2 (Muhtasim Ahmed)")
    features_list = [
        "360-Degree Interactive 3D Canvas: Renders procedural high-detail 3D furniture models directly in the browser using Three.js.",
        "Mouse & Touch Orbit Controls: Click and drag (or touch swipe) to rotate the furniture item across horizontal and vertical axes.",
        "360° Auto-Rotation Toggle: Button to toggle automated smooth background rotation.",
        "Step Rotation Controls: Counter-clockwise (↺ 15°) and clockwise (↻ 15°) rotation buttons.",
        "Zoom & Elevation Sliders: Sliders to adjust camera distance (2.5 to 6.5) and height elevation (-0.8m to +0.8m).",
        "Texture Material Swatches: Live material variant switching (TAN Autumn Leather #8C5A2B, FOREST Sage Green #435B4D, EBONY Black #2B2B2D, and WIREFRAME mode).",
        "Metadata Overlay: Displays POLYGONS: 124.2k, LOD LEVEL: ULTRA, and Archival Series № 422 matching Screenshot 2."
    ]
    for feat in features_list:
        p_feat = doc.add_paragraph(style='List Bullet')
        r_feat = p_feat.add_run(feat)
        r_feat.font.size = Pt(10)

    add_heading_2("Baseline & Integration Hooks")
    baseline_list = [
        "User Auth & Roles: Sign In / Registration modal with role selection (Buyer, Seller, Courier).",
        "Marketplace Catalog: Furniture catalog with search bar, category filtering (Chairs, Sofas, Tables), and AI condition grade filters (EXCELLENT, GOOD, FAIR).",
        "Product Details Page: Hero page matching Screenshot 1 with prominent launch button, gallery thumbnails, and craftsmanship details.",
        "Room Planner Hook (Module 2 F6): Spatial placement preview matching Screenshot 3 with Height Adjust (0.0m) and Rotation 142° dial.",
        "Escrow Cart Hook (Module 1 F13 / Module 2 F14): Shopping cart drawer displaying escrow protection status."
    ]
    for base in baseline_list:
        p_base = doc.add_paragraph(style='List Bullet')
        r_base = p_base.add_run(base)
        r_base.font.size = Pt(10)

    # Section 5
    add_heading_1("5. How to Run & Test the Application")
    add_heading_2("Step 1: Start the Development Server")
    p_step1 = doc.add_paragraph(
        "Open a terminal inside the project directory (c:\\Users\\PC\\Desktop\\CSE471_Project) and execute:\n"
        "npm run dev\n"
        "Then open http://localhost:5173 in your web browser."
    )
    p_step1.runs[0].font.size = Pt(10)

    add_heading_2("Step 2: Test Feature 2 (360° Interactive 3D Canvas)")
    test_steps = [
        "Select 'Mid-Century Modern Tan Leather Lounge Chair' ($450) from the Marketplace homepage.",
        "On the Product Details page, verify the '3D Model Available' badge and click 'LAUNCH INTERACTIVE 3D INSPECTOR' (matching Screenshot 1).",
        "Inside the 3D inspector modal (matching Screenshot 2):",
        "  • Click and drag mouse (or swipe touch screen) to rotate 360° around the furniture model.",
        "  • Click the '3D' button to turn auto-rotation on or off.",
        "  • Use ↺ and ↻ buttons for step rotation.",
        "  • Drag the zoom slider to move closer or further away.",
        "  • Drag the elevation slider to move the camera up or down.",
        "  • Click TAN, FOREST, or EBONY swatches to observe real-time material texture changes in WebGL.",
        "  • Click WIREFRAME to toggle WebGL polygon wireframe rendering.",
        "  • Verify POLYGONS: 124.2k and LOD LEVEL: ULTRA overlay."
    ]
    for ts in test_steps:
        p_ts = doc.add_paragraph(style='List Bullet')
        r_ts = p_ts.add_run(ts)
        r_ts.font.size = Pt(10)

    doc.save("c:\\Users\\PC\\Desktop\\CSE471_Project\\PROJECT_DOCUMENTATION.docx")
    print("Successfully generated PROJECT_DOCUMENTATION.docx")

if __name__ == "__main__":
    create_documentation()
